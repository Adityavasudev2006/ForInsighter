from __future__ import annotations

import asyncio
import hashlib
import json
import os
import uuid
import base64
import re
from pathlib import Path

from fastapi import APIRouter, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
import httpx
import pandas as pd
from sqlalchemy import desc, select

from agents.orchestrator import Orchestrator
from models.database import Document, SessionLocal
from models.schemas import DocumentResponse, DocumentStatus, DocumentSummary, EntityData, ExtractedQuestion, UploadLinkRequest, ViewManifest
from services.pdf_service import generate_highlighted_pdf, generate_pdf_derivative, get_derivative_pdf_path, merge_pdfs
from tasks.celery_tasks import process_document_task


router = APIRouter(prefix="/documents", tags=["documents"])
orchestrator = Orchestrator()

SUPPORTED_UPLOAD_SUFFIXES = {".pdf", ".xlsx", ".xls", ".csv", ".docx", ".png", ".jpg", ".jpeg", ".webp"}


def _file_type_from_suffix(suffix: str) -> str:
    if suffix == ".pdf":
        return "pdf"
    if suffix in {".xlsx", ".xls"}:
        return "excel"
    if suffix == ".csv":
        return "csv"
    if suffix == ".docx":
        return "docx"
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        return "image"
    return "text"


def _to_response(doc: Document) -> DocumentResponse:
    summary = DocumentSummary.model_validate(json.loads(doc.summary_json)) if doc.summary_json else None
    questions = [ExtractedQuestion.model_validate(q) for q in json.loads(doc.questions_json)] if doc.questions_json else None
    entities = EntityData.model_validate(json.loads(doc.entities_json)) if doc.entities_json else None
    view_manifest = ViewManifest.model_validate(json.loads(doc.view_manifest_json)) if doc.view_manifest_json else None
    if view_manifest is None:
        if doc.file_type == "pdf":
            view_manifest = ViewManifest(mode="pdf", source_mode="pdf")
        elif doc.file_type in {"excel", "csv"}:
            view_manifest = ViewManifest(mode="table", source_mode="native")
        elif doc.file_type in {"text", "docx"}:
            view_manifest = ViewManifest(mode="text", source_mode="native")
        elif doc.file_type == "image":
            view_manifest = ViewManifest(mode="image", source_mode="native")
    if view_manifest is not None:
        if not view_manifest.view_url:
            view_manifest.view_url = f"/api/documents/{doc.id}/native-view"
        if not view_manifest.highlight_url:
            view_manifest.highlight_url = f"/api/documents/{doc.id}/native-highlight"
    return DocumentResponse(
        id=doc.id,
        filename=doc.filename,
        file_type=doc.file_type,
        doc_type=doc.doc_type,
        status=DocumentStatus(doc.status),
        created_at=doc.created_at,
        summary=summary,
        questions=questions,
        entities=entities,
        processing_error=doc.processing_error,
        view_manifest=view_manifest,
    )


async def _process_document_in_process(doc_id: str, file_path: str, file_type: str, mode: str, llm_options: dict | None = None) -> None:
    async with SessionLocal() as session:
        doc = await session.get(Document, doc_id)
        if doc is None:
            return
        doc.status = "processing"
        doc.processing_error = None
        await session.commit()
    try:
        result = await orchestrator.process_document(doc_id, file_path, file_type, mode=mode, llm_options=llm_options or {})
        async with SessionLocal() as session:
            doc = await session.get(Document, doc_id)
            if doc is None:
                return
            summary_payload = dict(result["summary"])
            if isinstance(result.get("profile"), dict):
                summary_payload.update(result["profile"])
            doc.summary_json = json.dumps(summary_payload)
            doc.questions_json = json.dumps(result["questions"])
            doc.entities_json = json.dumps(result["entities"])
            if result.get("profile") is not None:
                doc.profile_json = json.dumps(result["profile"])
            if result.get("view_manifest") is not None:
                doc.view_manifest_json = json.dumps(result["view_manifest"])
            doc.doc_type = result["doc_type"]
            doc.status = "done"
            doc.processing_error = None
            await session.commit()
    except Exception as exc:
        async with SessionLocal() as session:
            doc = await session.get(Document, doc_id)
            if doc is None:
                return
            doc.status = "failed"
            doc.processing_error = str(exc)[:2000]
            await session.commit()


@router.post("/upload", response_model=list[DocumentResponse])
async def upload_documents(
    files: list[UploadFile] = File(...),
    mode: str = Query(default=os.getenv("LLM_MODE", "local")),
    api_provider: str | None = Query(default=None),
    api_key: str | None = Query(default=None),
    ollama_base_url: str | None = Query(default=None),
    ollama_model: str | None = Query(default=None),
):
    upload_dir = Path(os.getenv("UPLOAD_DIR", "./uploads"))
    upload_dir.mkdir(parents=True, exist_ok=True)
    output: list[DocumentResponse] = []

    # If multiple files are uploaded together, merge into one combined document tile.
    can_merge = len(files) > 1
    if can_merge:
        for file in files:
            suffix = Path(file.filename or "").suffix.lower()
            if _file_type_from_suffix(suffix) == "image":
                can_merge = False
                break

    if can_merge:
        if len(files) > 50:
            raise HTTPException(status_code=400, detail="Max 50 files per combined upload")

        saved_inputs: list[tuple[str, str, str]] = []  # (orig_name, file_type, saved_path)
        for file in files:
            name = file.filename or ""
            suffix = Path(name).suffix.lower()
            if suffix not in SUPPORTED_UPLOAD_SUFFIXES:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {name}")
            file_type = _file_type_from_suffix(suffix)
            content = await file.read()
            new_name = f"{uuid.uuid4()}{suffix}"
            saved_path = upload_dir / new_name
            saved_path.write_bytes(content)
            saved_inputs.append((name, file_type, str(saved_path)))

        # Convert non-PDF inputs to PDFs for merging.
        pdf_paths: list[str] = []
        tmp_derivatives: list[Path] = []
        for _, ft, path in saved_inputs:
            if ft == "pdf":
                pdf_paths.append(path)
            else:
                tmp_id = str(uuid.uuid4())
                derived = get_derivative_pdf_path(tmp_id)
                generate_pdf_derivative(tmp_id, path, ft)
                pdf_paths.append(str(derived))
                tmp_derivatives.append(derived)

        merged_name = f"{uuid.uuid4()}.pdf"
        merged_path = upload_dir / merged_name
        merge_pdfs(str(merged_path), pdf_paths)

        # Clean up input files and temp derivatives (the merged PDF becomes the stored artifact).
        for _, _, p in saved_inputs:
            try:
                Path(p).unlink(missing_ok=True)
            except Exception:
                pass
        for d in tmp_derivatives:
            try:
                d.unlink(missing_ok=True)
            except Exception:
                pass

        combined_filename = f"Merged ({len(files)} files)"
        combined_hash = hashlib.sha256(merged_path.read_bytes()).hexdigest()
        async with SessionLocal() as session:
            doc = Document(
                filename=combined_filename,
                file_path=str(merged_path),
                file_type="pdf",
                status="queued",
                content_hash=combined_hash,
                processing_error=None,
            )
            session.add(doc)
            await session.commit()
            await session.refresh(doc)
            llm_options = {
                "mode": mode,
                "api_provider": api_provider,
                "api_key": api_key,
                "ollama_base_url": ollama_base_url,
                "ollama_model": ollama_model,
            }
            try:
                process_document_task.delay(doc.id, doc.file_path, doc.file_type, mode, llm_options)
            except Exception:
                asyncio.create_task(_process_document_in_process(doc.id, doc.file_path, doc.file_type, mode, llm_options))
            output.append(_to_response(doc))
        return output

    for file in files:
        name = file.filename or ""
        suffix = Path(name).suffix.lower()
        if suffix not in SUPPORTED_UPLOAD_SUFFIXES:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {name}")
        file_type = _file_type_from_suffix(suffix)
        content = await file.read()
        content_hash = hashlib.sha256(content).hexdigest()
        new_name = f"{uuid.uuid4()}{suffix}"
        saved_path = upload_dir / new_name
        saved_path.write_bytes(content)
        async with SessionLocal() as session:
            doc = Document(
                filename=name,
                file_path=str(saved_path),
                file_type=file_type,
                status="queued",
                content_hash=content_hash,
                processing_error=None,
            )
            session.add(doc)
            await session.commit()
            await session.refresh(doc)
            llm_options = {
                "mode": mode,
                "api_provider": api_provider,
                "api_key": api_key,
                "ollama_base_url": ollama_base_url,
                "ollama_model": ollama_model,
            }
            try:
                process_document_task.delay(doc.id, doc.file_path, doc.file_type, mode, llm_options)
            except Exception:
                asyncio.create_task(_process_document_in_process(doc.id, doc.file_path, doc.file_type, mode, llm_options))
            output.append(_to_response(doc))
    return output


def _drive_file_id(url: str) -> str | None:
    # Common formats:
    # - https://drive.google.com/file/d/<id>/view?usp=sharing
    # - https://drive.google.com/open?id=<id>
    m = re.search(r"/file/d/([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    m = re.search(r"[?&]id=([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    return None


def _extract_google_form_questions(html: str) -> list[str]:
    # Best-effort extraction from FB_PUBLIC_LOAD_DATA_ array.
    m = re.search(r"FB_PUBLIC_LOAD_DATA_\s*=\s*(\[[\s\S]*?\]);", html)
    if not m:
        return []
    try:
        data = json.loads(m.group(1))
    except Exception:
        return []

    questions: list[str] = []

    def walk(node: object):
        if isinstance(node, str):
            t = " ".join(node.split())
            if 5 <= len(t) <= 200 and (t.endswith("?") or t.lower().startswith(("what", "why", "how", "when", "where", "who", "which", "please"))):
                questions.append(t)
            return
        if isinstance(node, list):
            for item in node:
                walk(item)
        if isinstance(node, dict):
            for v in node.values():
                walk(v)

    walk(data)
    # De-dupe preserving order.
    out: list[str] = []
    seen: set[str] = set()
    for q in questions:
        norm = q.lower()
        if norm not in seen:
            seen.add(norm)
            out.append(q)
    return out


@router.post("/upload-link", response_model=DocumentResponse)
async def upload_document_from_link(
    payload: UploadLinkRequest,
    mode: str = Query(default=os.getenv("LLM_MODE", "local")),
    api_provider: str | None = Query(default=None),
    api_key: str | None = Query(default=None),
    ollama_base_url: str | None = Query(default=None),
    ollama_model: str | None = Query(default=None),
):
    url = (payload.url or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="url is required")

    upload_dir = Path(os.getenv("UPLOAD_DIR", "./uploads"))
    upload_dir.mkdir(parents=True, exist_ok=True)

    llm_options = {
        "mode": mode,
        "api_provider": api_provider,
        "api_key": api_key,
        "ollama_base_url": ollama_base_url,
        "ollama_model": ollama_model,
    }

    # Resolve short links / redirects first (forms.gle, drive share variants).
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30) as client:
            head = await client.get(url)
            final_url = str(head.url)
            body_text = head.text
    except Exception:
        final_url = url
        body_text = ""

    url = final_url or url

    # Google Drive PDF share
    drive_id = _drive_file_id(url)
    if drive_id:
        async def fetch_drive_pdf_bytes() -> bytes:
            dl = f"https://drive.google.com/uc?export=download&id={drive_id}"
            async with httpx.AsyncClient(follow_redirects=True, timeout=60) as client:
                resp = await client.get(dl)
                if resp.status_code >= 400:
                    raise HTTPException(status_code=400, detail="Unable to fetch Google Drive file")
                ct = (resp.headers.get("content-type") or "").lower()
                if "application/pdf" in ct:
                    return resp.content
                # Google Drive sometimes returns an HTML interstitial with a confirm token.
                html = resp.text
                m = re.search(r"confirm=([0-9A-Za-z_]+)", html)
                if not m:
                    # If still not a PDF, treat as failure to avoid saving HTML as PDF.
                    raise HTTPException(status_code=400, detail="Google Drive did not return a PDF (permission or confirm page)")
                token = m.group(1)
                resp2 = await client.get(f"https://drive.google.com/uc?export=download&confirm={token}&id={drive_id}")
                if resp2.status_code >= 400:
                    raise HTTPException(status_code=400, detail="Unable to fetch Google Drive file (confirm)")
                ct2 = (resp2.headers.get("content-type") or "").lower()
                if "application/pdf" not in ct2:
                    raise HTTPException(status_code=400, detail="Google Drive did not return a PDF (confirm)")
                return resp2.content

        content = await fetch_drive_pdf_bytes()
        content_hash = hashlib.sha256(content).hexdigest()
        new_name = f"{uuid.uuid4()}.pdf"
        saved_path = upload_dir / new_name
        saved_path.write_bytes(content)
        async with SessionLocal() as session:
            doc = Document(
                filename=f"gdrive_{drive_id}.pdf",
                file_path=str(saved_path),
                file_type="pdf",
                status="queued",
                content_hash=content_hash,
                processing_error=None,
            )
            session.add(doc)
            await session.commit()
            await session.refresh(doc)
            try:
                process_document_task.delay(doc.id, doc.file_path, doc.file_type, mode, llm_options)
            except Exception:
                asyncio.create_task(_process_document_in_process(doc.id, doc.file_path, doc.file_type, mode, llm_options))
            return _to_response(doc)

    # Public Google Form (questions extraction)
    if "docs.google.com/forms" in url:
        html = body_text
        if not html:
            async with httpx.AsyncClient(follow_redirects=True, timeout=60) as client:
                resp = await client.get(url)
                if resp.status_code >= 400:
                    raise HTTPException(status_code=400, detail="Unable to fetch Google Form")
                html = resp.text
        questions = _extract_google_form_questions(html)
        if not questions:
            raise HTTPException(status_code=400, detail="Unable to extract questions from the Google Form")
        content_text = "Google Form Questions:\n\n" + "\n".join([f"- {q}" for q in questions]) + "\n"
        content_hash = hashlib.sha256(content_text.encode("utf-8")).hexdigest()
        new_name = f"{uuid.uuid4()}.txt"
        saved_path = upload_dir / new_name
        saved_path.write_text(content_text, encoding="utf-8")
        async with SessionLocal() as session:
            doc = Document(
                filename="google_form.txt",
                file_path=str(saved_path),
                file_type="text",
                status="queued",
                content_hash=content_hash,
                processing_error=None,
            )
            session.add(doc)
            await session.commit()
            await session.refresh(doc)
            try:
                process_document_task.delay(doc.id, doc.file_path, doc.file_type, mode, llm_options)
            except Exception:
                asyncio.create_task(_process_document_in_process(doc.id, doc.file_path, doc.file_type, mode, llm_options))
            return _to_response(doc)

    raise HTTPException(status_code=400, detail="Unsupported link. Provide a Google Drive PDF share link or a public Google Form link.")


@router.get("", response_model=list[DocumentResponse])
async def get_documents():
    async with SessionLocal() as session:
        result = await session.execute(select(Document).order_by(desc(Document.created_at)))
        docs = result.scalars().all()
    return [_to_response(doc) for doc in docs]


@router.get("/{doc_id}", response_model=DocumentResponse)
async def get_document(doc_id: str):
    async with SessionLocal() as session:
        doc = await session.get(Document, doc_id)
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")
    return _to_response(doc)


@router.delete("/{doc_id}")
async def delete_document(doc_id: str):
    async with SessionLocal() as session:
        doc = await session.get(Document, doc_id)
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")
        file_path = doc.file_path
        await session.delete(doc)
        await session.commit()
    if file_path and Path(file_path).exists():
        Path(file_path).unlink(missing_ok=True)
    # Delete derivative PDF if it exists.
    try:
        get_derivative_pdf_path(doc_id).unlink(missing_ok=True)
    except Exception:
        pass
    # Clear QA cache keys for this document (best-effort).
    try:
        keys = await orchestrator.cache_service.smembers(f"qa_keys:{doc_id}")
        await orchestrator.cache_service.delete_many(list(keys) + [f"qa_keys:{doc_id}"])
    except Exception:
        pass
    orchestrator.chroma_service.delete_collection(doc_id)
    return {"status": "deleted"}


@router.get("/{doc_id}/view")
async def view_document_pdf(doc_id: str):
    """
    Returns a PDF suitable for read-only viewing in the frontend.
    - PDF uploads: serves the original PDF.
    - Excel uploads: serves a generated PDF derivative (cached on disk).
    """
    async with SessionLocal() as session:
        doc = await session.get(Document, doc_id)
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")
        file_path = doc.file_path
        file_type = doc.file_type

    if file_type == "pdf":
        if not file_path or not Path(file_path).exists():
            raise HTTPException(status_code=404, detail="File not found")
        return FileResponse(
            file_path,
            media_type="application/pdf",
            headers={"Content-Disposition": f'inline; filename="{doc_id}.pdf"'},
        )
    if file_type == "image":
        ext = Path(file_path).suffix.lower().lstrip(".") or "png"
        media = "image/jpeg" if ext in {"jpg", "jpeg"} else f"image/{ext}"
        return FileResponse(file_path, media_type=media, headers={"Content-Disposition": f'inline; filename="{doc_id}.{ext}"'})

    derivative = get_derivative_pdf_path(doc_id)
    if not derivative.exists():
        try:
            generate_pdf_derivative(doc_id, file_path, file_type)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Unable to generate view: {exc}")
    return FileResponse(
        str(derivative),
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{doc_id}.pdf"'},
    )


def _highlight_text_html(text: str, needles: list[str]) -> str:
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    for needle in needles[:8]:
        clean = (needle or "").strip()
        if len(clean) < 5:
            continue
        safe = clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        escaped = escaped.replace(safe, f"<mark>{safe}</mark>")
    return escaped


@router.get("/{doc_id}/native-view")
async def native_view_document(doc_id: str, sources_b64: str | None = Query(default=None, alias="sources")):
    async with SessionLocal() as session:
        doc = await session.get(Document, doc_id)
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")
        file_path = doc.file_path
        file_type = doc.file_type

    source_snippets: list[str] = []
    if sources_b64:
        try:
            raw = base64.urlsafe_b64decode(sources_b64.encode("utf-8") + b"===")
            payload = json.loads(raw.decode("utf-8"))
            if isinstance(payload, list):
                source_snippets = [str(item.get("snippet", "")) for item in payload if isinstance(item, dict)]
        except Exception:
            source_snippets = []

    if file_type in {"text", "docx"}:
        if file_type == "docx":
            from docx import Document as DocxDocument
            d = DocxDocument(file_path)
            raw_text = "\n".join([p.text for p in d.paragraphs if p.text.strip()])
        else:
            raw_text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        html_body = _highlight_text_html(raw_text, source_snippets).replace("\n", "<br/>")
        return HTMLResponse(f"<html><body style='font-family:Arial;padding:12px;line-height:1.5'>{html_body}</body></html>")

    if file_type in {"excel", "csv"}:
        try:
            if file_type == "csv":
                df = pd.read_csv(file_path)
                html_table = df.head(3000).to_html(index=False, escape=True)
            else:
                xls = pd.ExcelFile(file_path)
                blocks: list[str] = []
                for sheet_name in xls.sheet_names:
                    sdf = xls.parse(sheet_name=sheet_name)
                    blocks.append(f"<h3>{sheet_name}</h3>{sdf.head(1000).to_html(index=False, escape=True)}")
                html_table = "\n".join(blocks)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Unable to render table: {exc}")
        marked = _highlight_text_html(html_table, source_snippets)
        return HTMLResponse(f"<html><body style='font-family:Arial;padding:12px'>{marked}</body></html>")

    if file_type == "image":
        return HTMLResponse(
            "<html><body style='margin:0;background:#111;display:flex;align-items:center;justify-content:center'>"
            f"<img src='/api/documents/{doc_id}/view' style='max-width:100%;max-height:100vh;object-fit:contain' alt='uploaded image'/>"
            "</body></html>",
            headers={"Content-Type": "text/html"},
        )

    # fallback for pdf
    return HTMLResponse(
        "<html><body style='margin:0'>"
        f"<iframe src='/api/documents/{doc_id}/view' style='border:0;width:100%;height:100vh'></iframe>"
        "</body></html>"
    )


@router.get("/{doc_id}/native-highlight")
async def native_highlight_document(doc_id: str, sources_b64: str = Query(..., alias="sources")):
    return await native_view_document(doc_id, sources_b64=sources_b64)


@router.get("/{doc_id}/highlight")
async def view_document_pdf_highlighted(doc_id: str, sources_b64: str = Query(..., alias="sources")):
    """
    Returns a highlighted PDF for a specific answer, based on source snippets + page numbers.
    Frontend encodes sources JSON using base64url to avoid state on the server.
    """
    async with SessionLocal() as session:
        doc = await session.get(Document, doc_id)
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")
        file_path = doc.file_path
        file_type = doc.file_type

    try:
        raw = base64.urlsafe_b64decode(sources_b64.encode("utf-8") + b"===")
        sources = json.loads(raw.decode("utf-8"))
        if not isinstance(sources, list):
            raise ValueError("sources must be a list")
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid sources payload")

    # Choose the base PDF for highlighting.
    if file_type == "pdf":
        base_pdf = file_path
    else:
        derivative = get_derivative_pdf_path(doc_id)
        if not derivative.exists():
            try:
                generate_pdf_derivative(doc_id, file_path, file_type)
            except Exception as exc:
                raise HTTPException(status_code=500, detail=f"Unable to generate view: {exc}")
        base_pdf = str(derivative)

    try:
        out_path = generate_highlighted_pdf(base_pdf, doc_id=doc_id, sources=sources)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Unable to generate highlights: {exc}")
    return FileResponse(
        str(out_path),
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{doc_id}.pdf"'},
    )
